
'''
Code References:

When designing the different prototypes, the following websites were instrumental in understanding more about some of the technologies included:

Prototype 1 (API Calls and JSON Parsing):
towardsdatascience.com, How to parse JSON data with Python Pandas?, Ankit Goel, https://towardsdatascience.com/how-to-parse-json-data-with-python-pandas-f84fbd0b1025, 04/07/2020

Prototype 2 (MySQL Database and Hashing & Salting):
W3Schools, Python MySQL, https://www.w3schools.com/python/python_mysql_getstarted.asp,
Geeksforgeeks, How To Hash Passwords In Python, https://www.geeksforgeeks.org/how-to-hash-passwords-in-python/, 28/12/2022

Prototype 3 (Scheduling and Email Notifications):
YouTube, Scheduling Tasks Professionally in Python, NeuralNine, https://www.youtube.com/watch?v=yDPQfj4bZY8&t=317s, 23/12/2022
medium.com, How to send emails with attachments with Python by using Microsoft Outlook or Office365 SMTP server, Michael King, https://medium.com/@neonforge/how-to-send-emails-with-attachments-with-python-by-using-microsoft-outlook-or-office365-smtp-b20405c9e63a, 26/12/2019
GitHub, dasgoll, https://gist.github.com/dasgoll/2f27ec2b703fbcb49a273d550d7865f8
FreeCodeCamp, Send Emails Using Python, https://www.freecodecamp.org/news/send-emails-using-code-4fcea9df63f, 07/10/2016
StackOverflow, Python how do i send multiple files in the email. I can send 1 file but how to send more than 1, https://stackoverflow.com/questions/37204979/python-how-do-i-send-multiple-files-in-the-email-i-can-send-1-file-but-how-to-s
Youtube, Corey Schafer, How to Send Emails Using Python - Plain Text, Adding Attachments, HTML Emails, and More,  https://www.youtube.com/watch?v=JRCJ6RtE3xU&t=1716s&ab_channel=CoreySchafer, 18/03/2019

Prototype 4 (GUI Development and Data Mining [Exploratory Data Analysis] and the features mentioned above):
All previous prototypes references
YouTube, NeuralNine , Modern Graphical User Interfaces in Python , https://www.youtube.com/watch?v=iM3kjbbKHQU, 21/11/2022
GitHub, CustomTkinter, Tom Schimansky, https://github.com/TomSchimansky/CustomTkinter/wiki, 06/12/2022
StackOverflow, background function in Python, https://stackoverflow.com/questions/7168508/background-function-in-python
Geeksforgeeks , Working with Images – Python .docx Module, https://www.geeksforgeeks.org/working-with-images-python-docx-module/, 21/11/2022
Seaborn, https://seaborn.pydata.org/

Prototype 5 (Twitter Web Scraping and fixing some issues in the app [enhancing user experience, error handling, etc.] and the features mentioned above):
All previous prototypes references
YouTube, Scrape Twitter with 5 Lines of Code, Rob Mulla https://www.youtube.com/watch?v=PUMMCLrVn8A&ab_channel=RobMulla
StackOverflow , Rotate label text in seaborn factorplot , https://stackoverflow.com/questions/26540035/rotate-label-text-in-seaborn-factorplot
Towardsdatascience , How to Change the Size of Figures in Matplotlib, https://towardsdatascience.com/change-figure-size-matplotlib-11a409f39584, 31/08/2021
Stackabuse, Change Font Size in Matplotlib, David Landup, https://stackabuse.com/change-font-size-in-matplotlib/
W3Docs, When to use cla(), clf() or close() for clearing a plot in matplotlib?, https://www.w3docs.com/snippets/python/when-to-use-cla-clf-or-close-for-clearing-a-plot-in-matplotlib.html
'''


import time
import customtkinter
import mysql.connector
import hashlib
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from email.mime.text import MIMEText
from email.message import EmailMessage
import schedule
import time as t
import os
import requests
import pandas as pd
import docx
import matplotlib.pyplot as plt
import seaborn as sns
import snscrape.modules.twitter as snt
import threading

################ DATABASE and HASHING STUFF ###############

try:
    database = mysql.connector.connect(
        host="localhost",
        user="root",
        password="",
        database="fyp"
    )

    db_cursor = database.cursor()
except:
    print('Could not connect to the database, make sure you have the database and that it is up and running')


def insertValues(api):
    sql = f"INSERT INTO accounts (API_Key_Hash_Value) VALUES ('{api}')"
    db_cursor.execute(sql)
    database.commit()


def listKeys():
    db_cursor.execute("SELECT * FROM accounts")
    result = db_cursor.fetchall()
    return [x for x in result]


def updateKey(old, new):
    sql = f"UPDATE accounts SET API_Key_Hash_Value = '{new}' WHERE API_Key_Hash_Value = '{old}'"
    db_cursor.execute(sql)
    database.commit()


def hash_api(api):
    salt = "abc"
    hashed_api = api + salt
    hashed = hashlib.sha256(hashed_api.encode())
    return hashed.hexdigest()


################ API CALL FUNCTIONS ###############


def api_call(user_api, max_longitude, max_latitude, min_longitude, min_latitude):
    URL = f'https://services.marinetraffic.com/api/exportvessels/v:8/{user_api}/MINLAT:{min_latitude}/MAXLAT:{max_latitude}/MINLON:{min_longitude}/MAXLON:{max_longitude}/timespan:10/protocol:jsono/shiptype:7/msgtype:extended'
    request = requests.get(URL).text
    df = pd.read_json(request)
    df.to_csv('data_extended.csv', index=False)


################### THREADING #####################


def thread_target():
    while True:
        schedule.run_pending()
        t.sleep(1)


def thread():
    return threading.Thread(target=thread_target, name="Notification")


################ EMAIL FUNCTIONS ###############

def send_notification(sender, password, receiver, e_subject, e_body):
    message = MIMEMultipart()
    message['From'] = sender
    message['To'] = receiver
    message['Subject'] = e_subject
    body = e_body
    message.attach(MIMEText(body, 'plain'))

    a = []

    if eda_report_chechbox.get() == 'on' and web_scraping_checbox.get() == 'on':
        a = ['data_extended.csv', 'Diags.docx', 'tweets.csv']
    elif eda_report_chechbox.get() == 'on' and web_scraping_checbox.get() == 'off':
        a = ['data_extended.csv', 'Diags.docx']
    elif eda_report_chechbox.get() == 'off' and web_scraping_checbox.get() == 'on':
        a = ['data_extended.csv', 'tweets.csv']
    elif eda_report_chechbox.get() == 'off' and web_scraping_checbox.get() == 'off':
        a = ['data_extended.csv']

    for file in a:
        f = open(file, 'rb')
        file_name = os.path.basename(file)
        part = MIMEBase('application', "octet-stream", Name=file)
        part.set_payload((f.read()))
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', "attachment; filename= %s" % file_name)
        message.attach(part)

    server = smtplib.SMTP('smtp-mail.outlook.com', 587)
    server.starttls()
    server.login(sender, password)
    server.send_message(message)
    server.quit()


def visualization():
    df = pd.read_csv('data_extended.csv')

    plt.rcParams['font.size'] = '7'
    plt.gcf().set_size_inches(7, 6)
    sns.countplot(x="FLAG", data=df)
    plt.savefig('diag1.png')
    plt.clf()

    plt.xticks(rotation=30)
    sns.countplot(x="TYPE_NAME", data=df)
    plt.savefig('diag2.png')
    plt.clf()

    plt.xticks(rotation=30)
    sns.countplot(x="YEAR_BUILT", data=df)
    plt.savefig('diag3.png')
    plt.clf()

    sns.boxplot(data=df, y="TYPE_NAME", x="LENGTH")
    # plt.gcf().set_size_inches(6, 6)
    plt.yticks(rotation=60)
    plt.savefig('diag4.png')
    plt.clf()

    sns.boxplot(data=df, y="TYPE_NAME", x="WIDTH")
    plt.yticks(rotation=60)
    plt.savefig('diag5.png')
    plt.clf()

    sns.boxplot(data=df, y="TYPE_NAME", x="YEAR_BUILT")
    plt.yticks(rotation=60)
    plt.savefig('diag6.png')
    plt.clf()

    sns.lineplot(data=df, x="YEAR_BUILT", y="LENGTH")
    plt.savefig('diag7.png')
    plt.clf()

    sns.lineplot(data=df, x="YEAR_BUILT", y="WIDTH")
    plt.savefig('diag8.png')
    plt.clf()

    sns.lineplot(data=df, x="YEAR_BUILT", y="DRAUGHT")
    plt.savefig('diag9.png')
    plt.clf()

    doc = docx.Document()

    doc.add_heading('Diagrams', 0)

    doc.add_heading('Plot 1:', 3)
    doc.add_picture('diag1.png')

    doc.add_heading('Plot 2:', 3)
    doc.add_picture('diag2.png')

    doc.add_heading('Plot 3:', 3)
    doc.add_picture('diag3.png')

    doc.add_heading('Plot 4:', 3)
    doc.add_picture('diag4.png')

    doc.add_heading('Plot 5:', 3)
    doc.add_picture('diag5.png')

    doc.add_heading('Plot 6:', 3)
    doc.add_picture('diag6.png')

    doc.add_heading('Plot 7:', 3)
    doc.add_picture('diag7.png')

    doc.add_heading('Plot 8:', 3)
    doc.add_picture('diag8.png')

    doc.add_heading('Plot 9:', 3)
    doc.add_picture('diag9.png')

    doc.save('Diags.docx')


def dataset_filter(column, value):
    df = pd.read_csv('data_extended.csv')

    if len(df) > 0 and df[column].dtype == 'object' and value in df[column].unique():
        # print("Success")
        return True

    elif len(df) > 0 and df[column].dtype != 'object' and True in set(df[column] <= value):
        # print("Success")
        return True
    else:
        # print("Filter wasn't met")
        return False


def notification():
    print('Hello guys')


def time_range(sec):
    schedule.every(int(sec)).seconds.do(legit_not)
    thread().start()


def time_sched(clock):
    schedule.every().day.at(clock).do(legit_not)
    thread().start()


def read():
    df = pd.read_csv('data_extended.csv')
    return df['SHIPNAME']


def scrape(num_of_tweets):
    ships = []
    dates = []
    users = []
    links = []
    contents = []
    likes = []
    retweets = []

    for i in read():

        scraped = snt.TwitterSearchScraper('"' + i + '"')

        counter = 0

        for t in scraped.get_items():
            ships.append(i)
            dates.append(str(t.date))
            users.append(t.user.username)
            contents.append(t.rawContent)
            links.append(t.url)
            likes.append(t.likeCount)
            retweets.append(t.retweetCount)

            counter += 1
            if counter == num_of_tweets:
                break

    df = pd.DataFrame()
    df['Ship'] = ships
    df['Tweet_Date'] = dates
    df['Tweet_Auther'] = users
    df['Tweet_URL'] = links
    df['Tweet_Content'] = contents
    df['Tweet_Likes'] = likes
    df['Tweet_Retweets'] = retweets

    return df


################ GUI FUNCTIONS ###############


api_key = ''


def log_func():
    global api_key
    bool = False
    for i in listKeys():
        if hash_api(api_key_login_entry.get()) in i:
            bool = True
            api_key = api_key_login_entry.get()
            break
    if bool == True:
        second_window()
    else:
        popup_window('This API key is not in the database, signup first', 'Authentication popup')


def sign_func():
    bool = True
    for i in listKeys():
        if hash_api(api_key_signup_entry.get()) in i:
            bool = False
            break
    if bool == False:
        popup_window('API key is already in the database', 'Authentication popup')
    else:
        insertValues(hash_api(api_key_signup_entry.get()))
        popup_window('API key is added successfully', 'Authentication popup')


def update_func():
    global api_key
    bool = False
    for i in listKeys():
        if hash_api(api_key) in i:
            bool = True
            break
    print(bool)
    if bool == True:
        updateKey(hash_api(api_key), hash_api(new_api_entry.get()))
        api_key = new_api_entry.get()
        popup_window('api key changed to ' + api_key, 'API key change')


def hints_popup():
    hints = '''
    Hints:
    If you encounter any problems while using the app, please follow the following instructions :

    - Make sure you have your database up and running
    - Check that your API key is correct, if not, you can always sign up with a new one 
    or even update your API key in the Settings section
    - Check that you have credits in you MarineTraffic account so the app can use your
     API key to send API requests, if not then you will need to purchase more credits from Marine Traffic Website.
    - Check that you have the correct Max and Min Longitude and Latitude or you might find more or less ships than you expected.
    - Make sure you pressed the Submit Filter button before submitting your notification.
    - Make sure you pressed the Add Receiver Email To List button after entering a receiver email.
    - Make sure you pressed the Submit Notification button when you have supplied all required information.
    - Make sure the credentials for the sender email are correct and make sure that it is an Outlook email.
    - Make sure you enter the correct values for all the entries (eg. don't enter a digit in an entry expecting a string)
    - When you are specifying a column filter, make sure that the value you provide and the column you select are of the same type,
      i.e. if the column is numerical (eg. LENGTH) then the value should be numerical (eg. 200), and the same thing goes
      for categorical columns (eg. FLAG) and categorical values (eg. NL) 
    - You can always check the receiver emails you specified if you press the Check Receiver Emails List button.
    - When you filter for a value of a categorical column (eg. FLAG) then the 
    app will look for an exact match inside the data received from the API response, 
    but if you filter for a value of a numerical column (eg. LENGTH) then the app will 
    look for values equal to or less than the value you specified from the API response.

    Sample values for all the filters:

    Area_Max_Long: 	-4.13721
    Area_Max_Lat: 	44.24562
    Area_Min_Long: 	-6.7749
    Area_Min_Lat: 	38.20882
    TIME: 15:00
    TIME RANGE: 30
    FLAG: NL
    LENGTH: 100
    Number of tweets to scrape per ship: 1

    If you encounter any other problems that are not highlighted in here, 
    please feel free to post your issue in the github repo issues section and highlight the issue you are facing.   
    '''

    popup_window(hints, 'Hints popup')


def popup_window(message, title):
    window = customtkinter.CTkToplevel()
    window.geometry("850x670")
    window.title(title)

    label = customtkinter.CTkLabel(window, text=message)
    label.pack(side="top", fill="both", expand=True, padx=40, pady=40)


def switcher():
    if mode_switcher1.get() != 'on':
        customtkinter.set_appearance_mode('Light')
    else:
        customtkinter.set_appearance_mode('dark')


def switcher2():
    if mode_switcher2.get() != 'on':
        customtkinter.set_appearance_mode('Light')
    else:
        customtkinter.set_appearance_mode('dark')


receiver_emails = []


def test():
    print('Hello')


def legit_not():
    try:

        attr_list = [email_sender_entry.get(), sender_password_entry.get(), list(set(receiver_emails)),
                     email_subject_entry.get(),
                     email_text_area.get("0.0", "end")]

        # Important
        # API Request
        api_call(submit_filter()[0], submit_filter()[1], submit_filter()[2], submit_filter()[3], submit_filter()[4])

        df = pd.read_csv('data_extended.csv')

        # Filter
        if int(len(df)) > 0 and submit_filter()[8].isdigit():
            if dataset_filter(str(submit_filter()[7]), int(submit_filter()[8])) == True:
                print('Filter successfully met')
                prompt_text_area.insert(0.0, 'Filter successfully met\n')
                visualization()
                scrape(int(submit_filter()[9])).to_csv('tweets.csv')
                # Email Notification
                for i in range(len(attr_list[2])):
                    send_notification(attr_list[0], attr_list[1], attr_list[2][i], attr_list[3], attr_list[4])
                    print('Email ' + str(i + 1) + ' Sent Successfully')
                    prompt_text_area.insert(0.0, 'Email ' + str(i + 1) + ' Sent Successfully\n')
            else:
                print('Filter was not met')
                prompt_text_area.insert(0.0, 'Filter was not met\n')

            # Filter
        elif int(len(df)) > 0 and not submit_filter()[8].isdigit():
            if dataset_filter(str(submit_filter()[7]), str(submit_filter()[8])) == True:
                print('Filter successfully met')
                prompt_text_area.insert(0.0, 'Filter successfully met\n')
                visualization()
                scrape(int(submit_filter()[9])).to_csv('tweets.csv')
                # Email Notification
                for i in range(len(attr_list[2])):
                    send_notification(attr_list[0], attr_list[1], attr_list[2][i], attr_list[3], attr_list[4])
                    print('Email ' + str(i + 1) + ' Sent Successfully')
                    prompt_text_area.insert(0.0, 'Email ' + str(i + 1) + ' Sent Successfully\n')
            else:
                print('Filter was not met')
                prompt_text_area.insert(0.0, 'Filter was not met\n')
    except:
        popup_window('''An error took place, this error is linked to email credentials, column filtering value, or other issue,
        please close the app, rerun it and read the instructions in the hints popup window to avoid getting an error again''',
                     'Error')


def not_submit():
    try:
        not_submit_button.configure(state="disabled")
        submit_filter_button.configure(state="disabled")
        receiver_emails_button.configure(state="disabled")
        set_api_button.configure(state="disabled")

        float(submit_filter()[1])
        float(submit_filter()[2])
        float(submit_filter()[3])
        float(submit_filter()[4])
        int(submit_filter()[9])

        if submit_filter()[5] == 'TIME':
            time_sched(submit_filter()[6])

            print('Notification submitted')
            prompt_text_area.insert(0.0, 'Notification submitted\n')

        elif submit_filter()[5] == 'TIME RANGE':
            time_range(submit_filter()[6])
            print('Notification submitted')
            prompt_text_area.insert(0.0, 'Notification submitted\n')
    except:
        popup_window('''An error took place, This error is linked to the sea area coordinates, the time/time range value or num of scraped tweets value, 
        please close the app, rerun it and read the instructions in the hints popup window to avoid getting an error again''',
                     'Error')


def combobox_callback():
    print("combobox dropdown clicked:", columns_combobox.get())


def print_out():
    print("You typed", filter_entry.get())


def change_api():
    print('Your API Changed from', api_key, 'to', new_api_entry.get())


def add_receiver_to_list():
    receiver_emails.append(receiver_emails_entry.get())
    popup_window('Email was added successfully', 'Email adding popup')
    # print()


def check_receiver_emails():
    # print(receiver_emails)
    a = ''
    for i in set(receiver_emails):
        a += (i + '\n')
    popup_window(a, 'Receiver Emails')


def submit_filter():
    filters = []
    filters.append(api_key)
    filters.append(str(max_long_entry.get()))
    filters.append(str(max_lat_entry.get()))
    filters.append(str(min_long_entry.get()))
    filters.append(str(min_lat_entry.get()))
    filters.append(str(time_combobox.get()))
    filters.append(str(time_entry.get()))
    filters.append(str(columns_combobox.get()))
    filters.append(str(filter_entry.get()))
    filters.append(str(tweets_num_entry.get()))
    # print(filters)
    return (filters)


#######################################


def second_window():
    try:
        authentication_frame.destroy()
    except:
        print('Error')

    notifications_frame = customtkinter.CTkFrame(master=root)
    notifications_frame.pack(pady=20, padx=10, fill='both', expand=True, side='left')

    settings_frame = customtkinter.CTkFrame(master=root)
    settings_frame.pack(pady=20, padx=10, fill='both', expand=True, side='right')

    filters_frame = customtkinter.CTkFrame(master=root)
    filters_frame.pack(pady=20, padx=10, fill='both', expand=True, side='bottom')

    settings_label = customtkinter.CTkLabel(master=settings_frame, text='Settings', font=("Helvetica", 30, 'bold'))
    settings_label.pack(pady=20, padx=10)

    global mode_switcher2
    mode_switcher2 = customtkinter.StringVar(value="on")

    second_switch = customtkinter.CTkSwitch(master=settings_frame, text="Dark Mode", command=switcher,
                                            variable=mode_switcher1, onvalue="on", offvalue="off")
    second_switch.pack(pady=5, padx=10)

    hints_button = customtkinter.CTkButton(master=settings_frame, text='Hints', command=hints_popup)
    hints_button.pack(pady=5, padx=10)

    api_change_label = customtkinter.CTkLabel(master=settings_frame, text='Change API', font=("Helvetica", 14, 'bold'))
    api_change_label.pack(pady=5, padx=10)

    global new_api_entry
    new_api_entry = customtkinter.CTkEntry(master=settings_frame, placeholder_text='New API', width=520, show='*')
    new_api_entry.pack(pady=5, padx=10)

    global set_api_button
    set_api_button = customtkinter.CTkButton(master=settings_frame, text='Set New API', command=update_func)
    set_api_button.pack(pady=5, padx=10)

    prompt_label = customtkinter.CTkLabel(master=settings_frame, text='Prompt Area', font=("Helvetica", 14, 'bold'))
    prompt_label.pack(pady=5, padx=10)

    global prompt_text_area
    prompt_text_area = customtkinter.CTkTextbox(settings_frame, width=520, height=200)
    prompt_text_area.pack()
    prompt_text_area.configure(state="normal")

    ###

    create_notification_label = customtkinter.CTkLabel(master=notifications_frame, text='Creating Notifications',
                                                       font=("Helvetica", 30, 'bold'))
    create_notification_label.pack(pady=12, padx=10)

    set_sender_label = customtkinter.CTkLabel(master=notifications_frame, text='Set The Sender Email',
                                              font=("Helvetica", 14, 'bold'))
    set_sender_label.pack(pady=12, padx=10)

    global email_sender_entry
    email_sender_entry = customtkinter.CTkEntry(master=notifications_frame, placeholder_text='Sender Email', width=520)
    email_sender_entry.pack(pady=5, padx=10)

    global sender_password_entry
    sender_password_entry = customtkinter.CTkEntry(master=notifications_frame, placeholder_text='Sender Email Password',
                                                   width=520,
                                                   show='*')
    sender_password_entry.pack(pady=5, padx=10)

    email_body_label = customtkinter.CTkLabel(master=notifications_frame,
                                              text='Add Receiver Email To The List Of Receiver Emails',
                                              font=("Helvetica", 14, 'bold'))
    email_body_label.pack(pady=5, padx=10)

    global receiver_emails_entry
    receiver_emails_entry = customtkinter.CTkEntry(master=notifications_frame, placeholder_text='Receiver Email',
                                                   width=520)
    receiver_emails_entry.pack(pady=5, padx=10)

    global receiver_emails_button
    receiver_emails_button = customtkinter.CTkButton(master=notifications_frame, text='Add Receiver Email To List',
                                                     command=add_receiver_to_list)
    receiver_emails_button.pack(pady=5, padx=10)

    global email_subject_entry
    email_subject_entry = customtkinter.CTkEntry(master=notifications_frame, placeholder_text='Email Subject',
                                                 width=520)
    email_subject_entry.pack(pady=5, padx=10)

    email_body_label = customtkinter.CTkLabel(master=notifications_frame, text='Email Body:',
                                              font=("Helvetica", 14, 'bold'))
    email_body_label.pack(pady=5, padx=10)

    global email_text_area
    email_text_area = customtkinter.CTkTextbox(notifications_frame, width=520, height=100)
    email_text_area.pack()
    email_text_area.configure(state="normal")

    check_emails_button = customtkinter.CTkButton(master=notifications_frame, text='Check Receiver Emails List',
                                                  command=check_receiver_emails)
    check_emails_button.pack(pady=5, padx=10)

    global eda_report_chechbox
    eda_report_chechbox = customtkinter.CTkCheckBox(master=notifications_frame, text="EDA Report", onvalue="on",
                                                    offvalue="off")
    eda_report_chechbox.pack(padx=5, pady=10)

    global web_scraping_checbox
    web_scraping_checbox = customtkinter.CTkCheckBox(master=notifications_frame, text="Twitter Web Scraping Report",
                                                     onvalue="on",
                                                     offvalue="off")
    web_scraping_checbox.pack(padx=5, pady=10)

    global not_submit_button
    not_submit_button = customtkinter.CTkButton(master=notifications_frame, text='Submit Notification',
                                                command=not_submit)
    not_submit_button.pack(pady=5, padx=10)

    creating_filters_label = customtkinter.CTkLabel(master=filters_frame, text='Creating Filters',
                                                    font=("Helvetica", 30, 'bold'))
    creating_filters_label.pack(pady=12, padx=10)

    ###
    required_fields_label = customtkinter.CTkLabel(master=filters_frame, text='Required Fields (Fill All)',
                                                   font=("Helvetica", 14, 'bold'))
    required_fields_label.pack(pady=5, padx=10)

    global max_long_entry
    max_long_entry = customtkinter.CTkEntry(master=filters_frame, placeholder_text='Area_Max_Long', width=320)
    max_long_entry.pack(pady=2, padx=10)

    global max_lat_entry
    max_lat_entry = customtkinter.CTkEntry(master=filters_frame, placeholder_text='Area_Max_Lat', width=320)
    max_lat_entry.pack(pady=2, padx=10)

    global min_long_entry
    min_long_entry = customtkinter.CTkEntry(master=filters_frame, placeholder_text='Area_Min_Long', width=320)
    min_long_entry.pack(pady=2, padx=10)

    global min_lat_entry
    min_lat_entry = customtkinter.CTkEntry(master=filters_frame, placeholder_text='Area_Min_Lat', width=320)
    min_lat_entry.pack(pady=2, padx=10)

    ###

    global time_combobox
    time_combobox = customtkinter.CTkOptionMenu(master=filters_frame,
                                                values=['TIME', 'TIME RANGE'])
    time_combobox.pack(padx=20, pady=10)

    global time_entry
    time_entry = customtkinter.CTkEntry(master=filters_frame, placeholder_text='Time/Time Range', width=320)
    time_entry.pack(pady=2, padx=10)

    global columns_combobox
    columns_combobox = customtkinter.CTkOptionMenu(master=filters_frame,
                                                   values=['COURSE', 'STATUS', 'SHIP_NAME', 'SHIP_TYPE', 'FLAG',
                                                           'LENGTH',
                                                           'WIDTH', 'YEAR_BUILT', 'DESTINATION', 'DRAUGHT',
                                                           'DEADWEIGHT',
                                                           'GROSS_TONNAGE'],
                                                   )
    columns_combobox.pack(padx=20, pady=10)

    global filter_entry
    filter_entry = customtkinter.CTkEntry(master=filters_frame, placeholder_text='Fill a Value', width=320)
    filter_entry.pack(pady=12, padx=10)

    global tweets_num_entry
    tweets_num_entry = customtkinter.CTkEntry(master=filters_frame,
                                              placeholder_text='Number of tweets to scrape per ship', width=320)
    tweets_num_entry.pack(pady=12, padx=10)

    # filters_frame_button = customtkinter.CTkButton(master=filters_frame, text='Press', command=print_out)
    global submit_filter_button
    submit_filter_button = customtkinter.CTkButton(master=filters_frame, text='Submit Filters', command=submit_filter)
    submit_filter_button.pack(pady=12, padx=10)


################ GUI LAYOUT BUILDING ###############

customtkinter.set_appearance_mode('dark')
customtkinter.set_default_color_theme('green')

root = customtkinter.CTk()
root.title('Naval Scanner')

authentication_frame = customtkinter.CTkFrame(master=root)
authentication_frame.pack(pady=10, padx=10, fill='both', expand=True)

authentication_page_label = customtkinter.CTkLabel(master=authentication_frame, text='Authentication Page',
                                                   font=("Helvetica", 30, 'bold'))
authentication_page_label.pack(pady=12, padx=10)

login_field_label = customtkinter.CTkLabel(master=authentication_frame, text='Login Field',
                                           font=("Helvetica", 14, 'bold'))
login_field_label.pack(pady=12, padx=10)

api_key_login_entry = customtkinter.CTkEntry(master=authentication_frame, placeholder_text='API Key', show='*',
                                             width=520)
api_key_login_entry.pack(pady=12, padx=10)

login_button = customtkinter.CTkButton(master=authentication_frame, text='Login', command=log_func)
login_button.pack(pady=20, padx=10)

signup_label = customtkinter.CTkLabel(master=authentication_frame, text='Signup Field', font=("Helvetica", 14, 'bold'))
signup_label.pack(pady=12, padx=10)

api_key_signup_entry = customtkinter.CTkEntry(master=authentication_frame, placeholder_text='API Key', show='*',
                                              width=520)
api_key_signup_entry.pack(pady=12, padx=10)

signup_button = customtkinter.CTkButton(master=authentication_frame, text='Signup', command=sign_func)
signup_button.pack(pady=12, padx=10)

mode_switcher1 = customtkinter.StringVar(value="on")

first_switch = customtkinter.CTkSwitch(master=authentication_frame, text="Dark Mode", command=switcher,
                                       variable=mode_switcher1, onvalue="on", offvalue="off")
first_switch.pack(pady=12, padx=10)

root.mainloop()

